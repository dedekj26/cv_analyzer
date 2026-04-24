"""
DOCX Parser — Mengekstrak teks dari file DOCX (Microsoft Word).
Semua pemrosesan dilakukan di memori (tidak ada penulisan ke disk).
"""

import io
from loguru import logger
from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Mengekstrak teks dari file DOCX yang diberikan sebagai bytes.

    Args:
        file_bytes: Konten file DOCX dalam bentuk bytes.

    Returns:
        Teks yang diekstrak dari dokumen DOCX.

    Raises:
        ValueError: Jika file tidak dapat dibaca atau teks kosong.
    """
    try:
        docx_buffer = io.BytesIO(file_bytes)
        doc = Document(docx_buffer)

        all_text = []

        # Ekstrak teks dari paragraf
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)

        # Ekstrak teks dari tabel (jika ada)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    all_text.append(" | ".join(row_text))

        full_text = "\n\n".join(all_text).strip()

        if not full_text or len(full_text) < 50:
            raise ValueError(
                "Teks yang diekstrak dari DOCX terlalu sedikit. "
                "Pastikan file DOCX berisi konten yang memadai."
            )

        logger.info(
            f"DOCX berhasil diparsing: {len(doc.paragraphs)} paragraf, "
            f"{len(doc.tables)} tabel, {len(full_text)} karakter"
        )
        return full_text

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Gagal mengekstrak teks dari DOCX: {str(e)}")
        raise ValueError(
            f"Tidak dapat mengekstrak teks dari file DOCX. "
            f"Pastikan file tidak rusak. Detail: {str(e)}"
        )
