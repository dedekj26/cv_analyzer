"""
Script test untuk menguji endpoint /analyze dengan CV dummy.
Membuat file DOCX sederhana di memori lalu mengirimnya ke API.
"""
import io
import json
import urllib.request
from docx import Document

# 1. Buat CV dummy dalam format DOCX
doc = Document()
doc.add_heading("CURRICULUM VITAE", level=1)
doc.add_heading("Data Pribadi", level=2)
doc.add_paragraph("Nama: Ahmad Fauzi")
doc.add_paragraph("Email: ahmad.fauzi@email.com")
doc.add_paragraph("Telepon: +62 812-3456-7890")
doc.add_paragraph("LinkedIn: linkedin.com/in/ahmadfauzi")

doc.add_heading("Ringkasan Profesional", level=2)
doc.add_paragraph(
    "Software Engineer dengan pengalaman 3 tahun dalam pengembangan web. "
    "Berpengalaman dalam React, Node.js, dan Python. "
    "Mampu bekerja dalam tim dan memiliki kemampuan problem-solving yang baik."
)

doc.add_heading("Pengalaman Kerja", level=2)
doc.add_paragraph("Software Engineer - PT Teknologi Maju (2022 - Sekarang)")
doc.add_paragraph("- Mengembangkan aplikasi web menggunakan React dan Node.js")
doc.add_paragraph("- Melakukan code review untuk tim beranggotakan 5 orang")
doc.add_paragraph("- Berpartisipasi dalam daily standup dan sprint planning")

doc.add_paragraph("Junior Developer - CV Digital Kreatif (2021 - 2022)")
doc.add_paragraph("- Membantu pengembangan fitur frontend")
doc.add_paragraph("- Memperbaiki bug dan melakukan testing")

doc.add_heading("Pendidikan", level=2)
doc.add_paragraph("S1 Teknik Informatika - Universitas Indonesia (2017 - 2021)")
doc.add_paragraph("IPK: 3.45 / 4.00")

doc.add_heading("Keterampilan", level=2)
doc.add_paragraph("JavaScript, Python, React, Node.js, SQL, Git, Docker")

# Simpan ke buffer memori
buffer = io.BytesIO()
doc.save(buffer)
file_bytes = buffer.getvalue()

print(f"CV dummy dibuat: {len(file_bytes)} bytes")
print("Mengirim ke http://localhost:8001/analyze ...")
print("=" * 60)

# 2. Kirim ke API menggunakan multipart/form-data
boundary = "----TestBoundary123"
body = (
    f"--{boundary}\r\n"
    f'Content-Disposition: form-data; name="file"; filename="cv_test.docx"\r\n'
    f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n"
    f"\r\n"
).encode("utf-8") + file_bytes + f"\r\n--{boundary}--\r\n".encode("utf-8")

req = urllib.request.Request(
    "http://localhost:8001/analyze",
    data=body,
    headers={
        "Content-Type": f"multipart/form-data; boundary={boundary}",
    },
    method="POST",
)

try:
    with urllib.request.urlopen(req, timeout=60) as response:
        result = json.loads(response.read())
        print(json.dumps(result, indent=2, ensure_ascii=False))
        print("=" * 60)
        print(f"Skor: {result['score']}/100 - {result['label']}")
        print(f"Kekuatan: {len(result['strengths'])} poin")
        print(f"Kelemahan: {len(result['weaknesses'])} poin")
        print(f"Rekomendasi: {len(result['recommendations'])} poin")
        print(f"Waktu: {result['metadata']['processing_time_ms']}ms")
        print(f"Token: {result['metadata']['total_tokens']}")
except urllib.error.HTTPError as e:
    print(f"HTTP Error {e.code}: {e.read().decode()}")
except Exception as e:
    print(f"Error: {e}")
